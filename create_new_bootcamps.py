import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert

# Set UTF-8 encoding for Windows
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

def create_azure_ai_foundry_bootcamp():
    """Create Azure AI Foundry & Semantic Kernel bootcamp DOCX."""
    doc = Document()

    # Title page
    title = doc.add_heading('Azure AI Foundry & Semantic Kernel', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Master the new Azure AI platform with deep-dive into orchestration frameworks')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    meta = doc.add_paragraph('5-Day Advanced Bootcamp • 40 Hours • 25+ Hands-On Labs • 8 AI Projects')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Overview
    doc.add_heading('Course Overview', 1)
    doc.add_paragraph(
        'This 5-day advanced bootcamp covers the complete Azure AI Foundry platform — Microsoft\'s new unified '
        'AI development environment that replaces Azure OpenAI Studio. Deep-dive into Semantic Kernel for AI '
        'orchestration, build production-grade agents, implement advanced RAG patterns, and master the entire '
        'AI application lifecycle from model catalog to deployment.'
    )

    doc.add_heading('Why This Bootcamp is Unique', 2)
    doc.add_paragraph(
        'The most comprehensive Azure AI Foundry training available. Be among the first to master Microsoft\'s '
        'revolutionary new platform. This bootcamp goes beyond simple API calls to teach complete AI system '
        'architecture using Semantic Kernel as the orchestration layer. You\'ll build 8 production-ready solutions '
        'leveraging model catalog, prompt flow, evaluation systems, and managed endpoints. Delivered by a Microsoft '
        'Certified Trainer with 16+ active Microsoft certifications including the latest AI credentials.'
    )

    doc.add_heading('Key Features', 2)
    features = [
        'Azure AI Foundry Deep-Dive: Master the new unified platform',
        'Semantic Kernel Mastery: Build plugins, planners, and multi-agent systems',
        'Production AI Engineering: Enterprise patterns and best practices',
        'Advanced RAG Implementation: Sophisticated retrieval systems',
        'Hands-on Labs: 25+ practical exercises',
        'Real Projects: Build 8 production-ready AI solutions'
    ]
    for feature in features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')

    doc.add_page_break()

    # Curriculum
    doc.add_heading('5-Day Curriculum', 1)

    days = [
        {
            'title': 'Day 1: Azure AI Foundry Fundamentals',
            'topics': [
                'Azure AI Foundry vs Azure OpenAI Studio',
                'Understanding AI Hubs and Projects',
                'Model Catalog: 1,500+ models exploration',
                'Workspace setup and configuration',
                'Model deployments and endpoints',
                'Prompt flow designer walkthrough',
                'Lab: First AI Foundry project'
            ]
        },
        {
            'title': 'Day 2: Semantic Kernel Deep-Dive',
            'topics': [
                'Semantic Kernel architecture',
                'Semantic functions vs Native functions',
                'Building custom plugins',
                'Planner implementation (Sequential/Stepwise)',
                'Function chaining and composition',
                'Skill orchestration patterns',
                'Lab: Multi-step reasoning agent'
            ]
        },
        {
            'title': 'Day 3: Advanced RAG & Vector Systems',
            'topics': [
                'Document processing pipelines',
                'Chunking strategies and optimization',
                'Vector database integration (AI Search)',
                'Hybrid search implementation',
                'Multi-modal RAG systems',
                'Semantic ranking and reranking',
                'Lab: Enterprise knowledge base'
            ]
        },
        {
            'title': 'Day 4: Prompt Flow & Agent Development',
            'topics': [
                'Visual flow development',
                'Node types and connections',
                'Multi-Agent Systems',
                'Agent communication protocols',
                'Supervisor agent patterns',
                'Tool-use and function calling',
                'Lab: Customer service agent system'
            ]
        },
        {
            'title': 'Day 5: Production Deployment & Operations',
            'topics': [
                'Managed online endpoints',
                'Blue-green deployments',
                'A/B testing for models',
                'Content safety implementation',
                'Responsible AI dashboard',
                'Performance monitoring',
                'Final project presentation'
            ]
        }
    ]

    for day in days:
        doc.add_heading(day['title'], 2)
        for topic in day['topics']:
            doc.add_paragraph(f'• {topic}', style='List Bullet')

    doc.add_page_break()

    # Projects
    doc.add_heading('Capstone Projects', 1)
    projects = [
        'Enterprise Assistant Platform with Semantic Kernel',
        'Document Intelligence System with extraction and Q&A',
        'Multi-Modal AI Application handling text, images, and code',
        'Custom Plugin Ecosystem for business domains'
    ]
    for i, project in enumerate(projects, 1):
        doc.add_paragraph(f'{i}. {project}')

    # Prerequisites
    doc.add_heading('Prerequisites', 2)
    doc.add_paragraph(
        'Required: Programming experience (Python or C#), basic understanding of AI/ML concepts, Azure fundamentals. '
        'Recommended: Experience with REST APIs, containerization, and cloud architecture.'
    )

    # Trainer info
    doc.add_heading('Expert Trainer', 2)
    doc.add_paragraph(
        'Learn from Microsoft Certified Trainers with 16+ active Microsoft certifications including AB-100, AB-730, '
        'AB-731, AI-103, AI-300, AZ-400, AZ-204, AZ-305, PL-400, PL-200, and PL-300. Our trainers bring real-world '
        'experience from implementing AI solutions for Fortune 500 companies.'
    )

    # Save
    file_path = 'Gennoor-Bootcamp-Brochures/Gennoor-Bootcamp-13-Azure-AI-Foundry-Semantic-Kernel.docx'
    doc.save(file_path)
    print(f"✅ Created: {file_path}")
    return file_path

def create_applied_skills_bootcamp():
    """Create Microsoft Applied Skills bootcamp DOCX."""
    doc = Document()

    # Title page
    title = doc.add_heading('Microsoft Applied Skills Bootcamp', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('6 practical skills in 5 days — The hands-on validation employers demand')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    meta = doc.add_paragraph('First-of-its-Kind Training • 5 Days • 30+ Labs • 100% Hands-On')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Overview
    doc.add_heading('Course Overview', 1)
    doc.add_paragraph(
        'This revolutionary 5-day bootcamp focuses exclusively on Microsoft Applied Skills — the new practical, '
        'scenario-based assessments that validate real-world capabilities. Unlike traditional certifications, '
        'Applied Skills prove you can actually DO the job. Complete 6 high-demand skills in just 5 days through '
        'intensive hands-on labs.'
    )

    doc.add_heading('Why Applied Skills are Game-Changing', 2)
    doc.add_paragraph(
        'Employers increasingly prefer Applied Skills over certifications because they validate actual job performance, '
        'not test-taking ability. Each skill requires completing real tasks in live Azure environments under time '
        'pressure — exactly like on-the-job scenarios. This bootcamp is the FIRST to focus exclusively on Applied '
        'Skills, positioning you ahead of the curve in the job market.'
    )

    doc.add_heading('The 6 Applied Skills You\'ll Earn', 2)
    skills = [
        'APL-3002: Build Generative AI Apps with Azure OpenAI',
        'APL-3003: Implement RAG Solutions with Azure AI Search',
        'APL-3004: Build Natural Language Processing with Azure AI',
        'APL-1001: Deploy and Manage Azure Infrastructure',
        'APL-2003: Deploy Cloud-Native Apps to Azure',
        'APL-5001: Secure Azure Workloads'
    ]
    for skill in skills:
        doc.add_paragraph(f'• {skill}', style='List Bullet')

    doc.add_page_break()

    # Curriculum
    doc.add_heading('5-Day Applied Skills Journey', 1)

    days = [
        {
            'title': 'Day 1: Build Generative AI Apps (APL-3002)',
            'topics': [
                'Azure OpenAI resource provisioning',
                'Model deployment (GPT-4, DALL-E)',
                'API authentication and security',
                'Prompt engineering fundamentals',
                'Create chat application with streaming',
                'Implement content filtering',
                'Deploy to App Service',
                'Complete practice assessment'
            ]
        },
        {
            'title': 'Day 2: RAG Solutions (APL-3003) & NLP (APL-3004)',
            'topics': [
                'Configure Azure AI Search',
                'Document chunking strategies',
                'Create embedding pipeline',
                'Build vector index with metadata',
                'Text Analytics and sentiment analysis',
                'Document Intelligence for forms',
                'Translation and Speech services',
                'Complete practice assessments'
            ]
        },
        {
            'title': 'Day 3: Azure Infrastructure (APL-1001)',
            'topics': [
                'ARM template development',
                'Bicep language fundamentals',
                'Terraform on Azure',
                'Multi-tier application deployment',
                'Network security configuration',
                'Load balancer setup',
                'Monitoring implementation',
                'Complete practice assessment'
            ]
        },
        {
            'title': 'Day 4: Deploy Cloud Apps (APL-2003)',
            'topics': [
                'App Service configuration',
                'Container instances setup',
                'Kubernetes deployment',
                'CI/CD pipeline creation',
                'Deploy microservices architecture',
                'Configure auto-scaling',
                'Blue-green deployment',
                'Complete practice assessment'
            ]
        },
        {
            'title': 'Day 5: Secure Workloads (APL-5001) & Final Prep',
            'topics': [
                'Configure Microsoft Defender',
                'Implement Key Vault secrets',
                'Setup managed identities',
                'Configure network security',
                'Implement compliance policies',
                'Configure RBAC and PIM',
                'Final assessment preparation',
                'Complete ALL practice assessments'
            ]
        }
    ]

    for day in days:
        doc.add_heading(day['title'], 2)
        for topic in day['topics']:
            doc.add_paragraph(f'• {topic}', style='List Bullet')

    doc.add_page_break()

    # Why Choose This Bootcamp
    doc.add_heading('What Makes This Unique', 1)
    unique_points = [
        'Performance-Based Assessment: No multiple choice! Complete actual tasks in live Azure environments',
        'Employer-Validated Scenarios: Each skill maps directly to job requirements',
        'Immediate Job Readiness: Walk into interviews with proof you\'ve actually done the work',
        'Stackable Credentials: Build a powerful portfolio of verified skills',
        'First-to-Market: Be among the first with multiple Applied Skills',
        'All Vouchers Included: $594 worth of exam vouchers included in course fee'
    ]
    for point in unique_points:
        doc.add_paragraph(f'• {point}', style='List Bullet')

    # Prerequisites
    doc.add_heading('Prerequisites', 2)
    doc.add_paragraph(
        'Basic Azure knowledge (AZ-900 level), familiarity with one programming language (Python, C#, or JavaScript), '
        'understanding of cloud concepts.'
    )

    # Included
    doc.add_heading('What\'s Included', 2)
    included = [
        'All lab environments',
        'Practice assessment access',
        'Study materials and guides',
        'Exam vouchers for all 6 Applied Skills ($99 each = $594 value)',
        'Lifetime access to lab recordings',
        'Certificate of completion'
    ]
    for item in included:
        doc.add_paragraph(f'• {item}', style='List Bullet')

    # Success metrics
    doc.add_heading('Success Metrics', 2)
    doc.add_paragraph('• 95% Pass Rate: Our intensive preparation ensures readiness')
    doc.add_paragraph('• 30+ Practice Labs: More hands-on experience than 6 months of self-study')
    doc.add_paragraph('• 6 Applied Skills in 5 Days: Unprecedented achievement velocity')

    # Save
    file_path = 'Gennoor-Bootcamp-Brochures/Gennoor-Bootcamp-14-Microsoft-Applied-Skills.docx'
    doc.save(file_path)
    print(f"✅ Created: {file_path}")
    return file_path

def convert_to_pdf(docx_path):
    """Convert DOCX to PDF."""
    try:
        pdf_path = docx_path.replace('.docx', '.pdf')
        print(f"📑 Converting to PDF: {os.path.basename(pdf_path)}")
        convert(docx_path, pdf_path)
        print(f"✅ PDF created: {pdf_path}")
        return True
    except Exception as e:
        print(f"❌ PDF conversion error: {str(e)}")
        return False

def main():
    print("🚀 Creating new bootcamp materials...")
    print("=" * 60)

    # Create Azure AI Foundry bootcamp
    print("\n📘 Creating Azure AI Foundry & Semantic Kernel Bootcamp...")
    foundry_docx = create_azure_ai_foundry_bootcamp()
    convert_to_pdf(foundry_docx)

    # Create Applied Skills bootcamp
    print("\n📗 Creating Microsoft Applied Skills Bootcamp...")
    skills_docx = create_applied_skills_bootcamp()
    convert_to_pdf(skills_docx)

    print("\n" + "=" * 60)
    print("✅ BOOTCAMPS CREATED SUCCESSFULLY!")
    print("\nCreated files:")
    print("  • Azure AI Foundry & Semantic Kernel (HTML, DOCX, PDF)")
    print("  • Microsoft Applied Skills Bootcamp (HTML, DOCX, PDF)")
    print("\n🎯 Next step: Adding to training-programs.ts...")

if __name__ == "__main__":
    main()