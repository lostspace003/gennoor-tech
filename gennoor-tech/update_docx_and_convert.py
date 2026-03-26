import os
import sys
from docx import Document
from docx2pdf import convert
import glob

# Set UTF-8 encoding for Windows
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# Define the replacements
replacements = {
    # Certification code updates
    'AI-900': 'AI-901',
    'AI-102': 'AI-103',
    'DP-100': 'AI-300',

    # Certification name updates
    'Azure AI Engineer Associate': 'Azure AI App & Agent Developer Associate',
    'Azure Data Scientist Associate': 'Machine Learning Operations (MLOps) Engineer Associate',

    # Trainer certification count update
    '12+ active Microsoft certifications': '16+ active Microsoft certifications',
    '12+ Microsoft certifications': '16+ active Microsoft certifications',
}

# Additional specific replacements for certification lists
cert_list_replacements = {
    'AB-100, AI-102, AZ-400, AZ-204, AZ-305, DP-100, PL-400, PL-200, and PL-300':
    'AB-100, AB-730, AB-731, AI-103, AI-300, AZ-400, AZ-204, AZ-305, PL-400, PL-200, and PL-300',

    'including AI-102': 'including AI-103',
    'including DP-100': 'including AI-300',
}

def update_paragraph(paragraph):
    """Update text in a paragraph while preserving formatting."""
    text_updated = False

    # Check each run in the paragraph
    for run in paragraph.runs:
        original_text = run.text
        new_text = original_text

        # Apply all replacements
        for old_text, new_text_replacement in {**replacements, **cert_list_replacements}.items():
            if old_text in new_text:
                new_text = new_text.replace(old_text, new_text_replacement)
                text_updated = True

        # Update the run text if changed
        if new_text != original_text:
            run.text = new_text

    return text_updated

def update_table_cell(cell):
    """Update text in table cells."""
    text_updated = False
    for paragraph in cell.paragraphs:
        if update_paragraph(paragraph):
            text_updated = True
    return text_updated

def update_docx_file(file_path):
    """Update a single DOCX file with all replacements."""
    print(f"\n📄 Processing: {os.path.basename(file_path)}")

    try:
        # Open the document
        doc = Document(file_path)
        changes_made = False

        # Update all paragraphs
        for paragraph in doc.paragraphs:
            if update_paragraph(paragraph):
                changes_made = True

        # Update all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if update_table_cell(cell):
                        changes_made = True

        # Update headers and footers
        for section in doc.sections:
            # Update header
            if section.header:
                for paragraph in section.header.paragraphs:
                    if update_paragraph(paragraph):
                        changes_made = True

            # Update footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if update_paragraph(paragraph):
                        changes_made = True

        if changes_made:
            # Save the updated document
            doc.save(file_path)
            print(f"   ✅ Updated successfully")
            return True
        else:
            print(f"   ℹ️ No changes needed")
            return False

    except Exception as e:
        print(f"   ❌ Error: {str(e)}")
        return False

def convert_to_pdf(docx_path):
    """Convert DOCX to PDF."""
    try:
        # Get the PDF path (same name, different extension)
        pdf_path = docx_path.replace('.docx', '.pdf')

        print(f"   📑 Converting to PDF: {os.path.basename(pdf_path)}")

        # Convert the file
        convert(docx_path, pdf_path)

        print(f"   ✅ PDF created successfully")
        return True
    except Exception as e:
        print(f"   ❌ PDF conversion error: {str(e)}")
        return False

def main():
    print("🔄 Starting DOCX update and PDF conversion process...")
    print("=" * 60)

    # Directories to process
    directories = [
        'Gennoor-Bootcamp-Brochures',
        'Gennoor-Tech-Course-TOCs'
    ]

    total_updated = 0
    total_converted = 0
    files_to_convert = []

    for directory in directories:
        dir_path = os.path.join(os.getcwd(), directory)
        print(f"\n📁 Processing directory: {directory}")
        print("-" * 40)

        if not os.path.exists(dir_path):
            print(f"   ⚠️ Directory not found: {dir_path}")
            continue

        # Find all DOCX files
        docx_files = glob.glob(os.path.join(dir_path, "*.docx"))
        print(f"   Found {len(docx_files)} DOCX files")

        for docx_file in docx_files:
            if update_docx_file(docx_file):
                total_updated += 1
                files_to_convert.append(docx_file)
            else:
                # Still add to conversion list if PDF doesn't exist
                pdf_path = docx_file.replace('.docx', '.pdf')
                if not os.path.exists(pdf_path):
                    files_to_convert.append(docx_file)

    # Convert updated files to PDF
    if files_to_convert:
        print("\n" + "=" * 60)
        print("📑 Converting DOCX files to PDF...")
        print("-" * 40)

        for docx_file in files_to_convert:
            if convert_to_pdf(docx_file):
                total_converted += 1

    # Summary
    print("\n" + "=" * 60)
    print("✅ PROCESS COMPLETE!")
    print(f"   📝 DOCX files updated: {total_updated}")
    print(f"   📑 PDF files created/updated: {total_converted}")

    if total_updated > 0:
        print("\n📋 Updated certifications:")
        print("   • AI-900 → AI-901")
        print("   • AI-102 → AI-103")
        print("   • DP-100 → AI-300")
        print("   • Trainer certifications: 12+ → 16+")
        print("   • Added: AB-730, AB-731 to trainer list")

    print("\n🎯 All files are now updated with the latest Microsoft certifications!")

if __name__ == "__main__":
    main()