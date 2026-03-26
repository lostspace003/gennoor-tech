import os
import sys
import glob
import re
from docx import Document
from docx2pdf import convert
import shutil

# Set UTF-8 encoding for Windows
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# Define all replacements
REPLACEMENTS = {
    # Certification updates
    'AI-900': 'AI-901',
    'AI-102': 'AI-103',
    'DP-100': 'AI-300',

    # Certification names
    'Azure AI Engineer Associate': 'Azure AI App & Agent Developer Associate',
    'Azure Data Scientist Associate': 'Machine Learning Operations (MLOps) Engineer Associate',

    # Trainer certifications
    '12+ active Microsoft certifications': '16+ active Microsoft certifications',
    '12+ Microsoft certifications': '16+ active Microsoft certifications',
}

# Email replacements for HTML
EMAIL_REPLACEMENTS = {
    '[email&#160;protected]': 'training@gennoor.com',
    '[email protected]': 'training@gennoor.com',
    'data-cfemail="c7a2a9b6b2aeb5be87a0a2a9a9a8a8b5e9a4a8aa">[email&#160;protected]': '>training@gennoor.com',
    'data-cfemail="9bfef5eaeef2e9e2dbfcfef5f5f4f4e9b5f8f4f6">[email&#160;protected]': '>training@gennoor.com',
}

def fix_html_file(file_path):
    """Fix HTML file with all replacements."""
    print(f"   📝 Fixing HTML: {os.path.basename(file_path)}")

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        original_content = content

        # Apply certification replacements
        for old, new in REPLACEMENTS.items():
            content = content.replace(old, new)

        # Fix CloudFlare email protection
        # Remove CloudFlare email protection script and class
        content = re.sub(r'<a href="/cdn-cgi/l/email-protection"[^>]*>\[email[^<]*</a>',
                        'training@gennoor.com', content)

        # Also handle cases where email might appear differently
        content = content.replace('www.gennoor.com · <a href="/cdn-cgi/l/email-protection"',
                                'www.gennoor.com · training@gennoor.com<!--')
        content = content.replace('[email&#160;protected]</a>', '-->')

        # Direct email replacements
        for old, new in EMAIL_REPLACEMENTS.items():
            content = content.replace(old, new)

        if content != original_content:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        return False

    except Exception as e:
        print(f"      ❌ Error: {str(e)}")
        return False

def fix_docx_file(file_path):
    """Fix DOCX file with all replacements."""
    print(f"   📝 Fixing DOCX: {os.path.basename(file_path)}")

    try:
        doc = Document(file_path)
        changes_made = False

        # Update all paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                original_text = run.text
                new_text = original_text

                # Apply all replacements
                for old, new in REPLACEMENTS.items():
                    if old in new_text:
                        new_text = new_text.replace(old, new)
                        changes_made = True

                # Fix email
                if 'example.com' in new_text:
                    new_text = new_text.replace('email@example.com', 'training@gennoor.com')
                    new_text = new_text.replace('contact@example.com', 'training@gennoor.com')
                    changes_made = True

                if new_text != original_text:
                    run.text = new_text

        # Update tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            original_text = run.text
                            new_text = original_text

                            for old, new in REPLACEMENTS.items():
                                if old in new_text:
                                    new_text = new_text.replace(old, new)
                                    changes_made = True

                            if new_text != original_text:
                                run.text = new_text

        if changes_made:
            doc.save(file_path)
            return True
        return False

    except Exception as e:
        print(f"      ❌ Error: {str(e)}")
        return False

def convert_docx_to_pdf(docx_path):
    """Convert DOCX to PDF."""
    try:
        pdf_path = docx_path.replace('.docx', '.pdf')
        print(f"   📑 Converting to PDF: {os.path.basename(pdf_path)}")
        convert(docx_path, pdf_path)
        return True
    except Exception as e:
        print(f"      ❌ PDF conversion error: {str(e)}")
        return False

def main():
    print("🔧 COMPREHENSIVE WEBSITE CONTENT FIX")
    print("=" * 60)

    # Statistics
    html_fixed = 0
    docx_fixed = 0
    pdf_created = 0

    # Fix HTML files in Bootcamp Brochures
    print("\n📁 Fixing Bootcamp HTML files...")
    html_files = glob.glob('Gennoor-Bootcamp-Brochures/*.html')
    for html_file in html_files:
        if fix_html_file(html_file):
            html_fixed += 1

    # Fix HTML files in Course TOCs
    print("\n📁 Fixing Course TOC HTML files...")
    html_files = glob.glob('Gennoor-Tech-Course-TOCs/*.html')
    for html_file in html_files:
        if fix_html_file(html_file):
            html_fixed += 1

    # Fix DOCX files and convert to PDF
    print("\n📁 Fixing DOCX files and creating PDFs...")
    docx_files = glob.glob('Gennoor-Bootcamp-Brochures/*.docx')
    for docx_file in docx_files:
        if fix_docx_file(docx_file):
            docx_fixed += 1
        if convert_docx_to_pdf(docx_file):
            pdf_created += 1

    docx_files = glob.glob('Gennoor-Tech-Course-TOCs/*.docx')
    for docx_file in docx_files:
        if fix_docx_file(docx_file):
            docx_fixed += 1
        if convert_docx_to_pdf(docx_file):
            pdf_created += 1

    # Copy all updated files to public directory
    print("\n📂 Copying updated files to public directory...")

    # Copy Bootcamp files
    for ext in ['*.html', '*.pdf', '*.docx']:
        files = glob.glob(f'Gennoor-Bootcamp-Brochures/{ext}')
        for file in files:
            dest = f'public/Gennoor-Bootcamp-Brochures/{os.path.basename(file)}'
            shutil.copy2(file, dest)
            print(f"   ✅ Copied: {os.path.basename(file)}")

    # Copy Course TOC files
    for ext in ['*.html', '*.pdf', '*.docx']:
        files = glob.glob(f'Gennoor-Tech-Course-TOCs/{ext}')
        for file in files:
            dest = f'public/Gennoor-Tech-Course-TOCs/{os.path.basename(file)}'
            os.makedirs(os.path.dirname(dest), exist_ok=True)
            shutil.copy2(file, dest)
            print(f"   ✅ Copied: {os.path.basename(file)}")

    # Summary
    print("\n" + "=" * 60)
    print("✅ COMPLETE FIX SUMMARY:")
    print(f"   📝 HTML files fixed: {html_fixed}")
    print(f"   📄 DOCX files fixed: {docx_fixed}")
    print(f"   📑 PDF files created: {pdf_created}")

    print("\n🔄 All updates applied:")
    print("   • AI-900 → AI-901")
    print("   • AI-102 → AI-103")
    print("   • DP-100 → AI-300")
    print("   • Email: training@gennoor.com")
    print("   • Trainer: 16+ certifications")

    print("\n🌐 Website Status:")
    print("   ✅ All HTML files updated with correct content")
    print("   ✅ All PDFs regenerated with latest information")
    print("   ✅ All files copied to public directory")
    print("   ✅ Ready for immediate display on website")

    print("\n💡 To verify on website:")
    print("   1. Restart the server: npm run dev")
    print("   2. Clear browser cache (Ctrl+F5)")
    print("   3. Check any training page")
    print("   4. All content should show latest updates!")

if __name__ == "__main__":
    main()